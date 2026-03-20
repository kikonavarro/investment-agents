"""
Thesis Writer Agent — redacta la tesis de inversion a partir de la valoracion.
REQUIERE que se haya ejecutado run_analyst() previamente para el ticker.
Lee los archivos generados por la valoracion (JSON) para escribir la tesis.
Guarda la tesis en la misma carpeta de la valoracion.
"""

import json
from datetime import date
from pathlib import Path

from agents.base import call_agent
from agents.analyst import load_valuation, get_valuation_path, _clean_ticker
from tools.document_generator import save_thesis_markdown, save_thesis_docx
from config.prompts import THESIS_WRITER
from config import settings
from config.settings import VALUATIONS_DIR


def run_thesis_writer(analysis: dict | str) -> str:
    """
    Redacta la tesis completa. Acepta:
      - dict: resultado de run_analyst() (JSON de valoracion)
      - str: ticker (busca la valoracion existente en disco)

    Guarda la tesis en Markdown y DOCX dentro de la carpeta de valoracion.

    Raises:
        ValueError si no existe valoracion previa para el ticker.
    """
    # Resolver el input
    if isinstance(analysis, str):
        # Input es un ticker: cargar valoracion existente
        ticker = analysis.upper()
        analysis = load_valuation(ticker)
        if analysis is None:
            raise ValueError(
                f"No existe valoracion para {ticker}. "
                f"Ejecuta primero: python main.py --analyst {ticker}"
            )
    elif isinstance(analysis, dict):
        ticker = analysis.get("ticker", "UNKNOWN")
        # Si el dict no tiene datos completos (viene de pipeline viejo),
        # intentar cargar de disco
        if "historical_data" not in analysis and "latest_financials" not in analysis:
            loaded = load_valuation(ticker)
            if loaded:
                analysis = loaded
            else:
                raise ValueError(
                    f"Datos insuficientes para escribir tesis de {ticker}. "
                    f"Ejecuta primero: python main.py --analyst {ticker}"
                )
    else:
        raise ValueError("Input debe ser un ticker (str) o dict de valoracion")

    ticker = analysis.get("ticker", "UNKNOWN")
    company = analysis.get("company", "")
    today = date.today().strftime("%d/%m/%Y")
    prompt_with_date = THESIS_WRITER.replace("[fecha]", today)

    message = (
        f"Fecha de analisis: {today}\n"
        f"Datos completos de la valoracion:\n{json.dumps(analysis, ensure_ascii=False, indent=2)}"
    )

    # Modo data-only: devolver datos preparados sin llamar a la API
    if settings.DATA_ONLY_MODE:
        print(f"  [thesis_writer] Modo data-only: datos preparados para {ticker}")
        print(f"  [thesis_writer] Valoracion JSON: {VALUATIONS_DIR / _clean_ticker(ticker) / f'{_clean_ticker(ticker)}_valuation.json'}")
        return json.dumps(analysis, ensure_ascii=False, indent=2)

    print(f"  [thesis_writer] Redactando tesis para {ticker}...")
    thesis = call_agent(
        system_prompt=prompt_with_date,
        user_message=message,
        model_tier="deep",
        max_tokens=4000,
        json_output=False,
    )

    # Guardar en carpeta de valoracion
    folder_name = _clean_ticker(ticker)
    valuation_dir = VALUATIONS_DIR / folder_name

    if valuation_dir.exists():
        # Guardar en carpeta de valoracion
        md_path = valuation_dir / f"{folder_name}_tesis_inversion.md"
        md_path.write_text(thesis, encoding="utf-8")
        print(f"  [thesis] Tesis MD: {md_path}")

        # DOCX tambien en la carpeta de valoracion
        docx_path = valuation_dir / f"{folder_name}_tesis_inversion.docx"
        _save_docx_to_path(thesis, ticker, company, str(docx_path))
    else:
        # Fallback: guardar en data/analyses/
        save_thesis_markdown(thesis, ticker)
        save_thesis_docx(thesis, ticker, company)

    return thesis


def _save_docx_to_path(thesis_text: str, ticker: str, company: str, path: str):
    """Guarda DOCX en un path especifico."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Portada
        doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TESIS DE INVERSION")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run(f"{ticker}" + (f" - {company}" if company else ""))
        run2.font.size = Pt(16)
        run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

        doc.add_paragraph()
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(date.today().strftime("%d de %B de %Y")).font.size = Pt(12)

        doc.add_page_break()

        # Contenido
        for line in thesis_text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped == "" or stripped == "---":
                doc.add_paragraph()
            else:
                doc.add_paragraph(stripped)

        doc.save(path)
        print(f"  [thesis] Tesis DOCX: {path}")
    except Exception as e:
        print(f"  [thesis] Aviso: No se pudo generar DOCX: {e}")
