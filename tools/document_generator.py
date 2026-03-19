"""
Document Generator — exporta tesis y análisis a DOCX y Markdown.
Sin dependencias de LaTeX. python-docx para Word, pathlib para Markdown.
"""
import re
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import ANALYSES_DIR, REPORTS_DIR


def save_thesis_markdown(thesis_text: str, ticker: str) -> Path:
    """
    Guarda la tesis en Markdown en data/analyses/.

    Returns:
        Path del archivo creado.
    """
    ANALYSES_DIR.mkdir(parents=True, exist_ok=True)
    today = date.today().strftime("%Y%m%d")
    path = ANALYSES_DIR / f"{today}_{ticker}_tesis.md"
    path.write_text(thesis_text, encoding="utf-8")
    print(f"  [doc] Tesis guardada: {path}")
    return path


def save_thesis_docx(thesis_text: str, ticker: str, company_name: str = "") -> Path:
    """
    Convierte la tesis Markdown a DOCX con formato profesional.

    Returns:
        Path del archivo .docx creado.
    """
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    today = date.today().strftime("%Y%m%d")
    path = REPORTS_DIR / f"{today}_{ticker}_tesis.docx"

    doc = _create_styled_document()

    # Portada
    _add_cover(doc, ticker, company_name, today)

    # Convertir Markdown a párrafos con formato
    _markdown_to_docx(doc, thesis_text)

    # Pie de página con disclaimer
    _add_footer(doc)

    doc.save(path)
    print(f"  [doc] DOCX guardado: {path}")
    return path


def save_analysis_json(analysis: dict, ticker: str) -> Path:
    """Guarda el JSON del analyst en data/analyses/."""
    import json
    ANALYSES_DIR.mkdir(parents=True, exist_ok=True)
    today = date.today().strftime("%Y%m%d")
    path = ANALYSES_DIR / f"{today}_{ticker}_analysis.json"
    path.write_text(json.dumps(analysis, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  [doc] Análisis guardado: {path}")
    return path


def save_screener_report(screener_result: dict) -> Path:
    """Genera un informe Markdown del screener."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    today = date.today().strftime("%Y%m%d")
    path = REPORTS_DIR / f"{today}_screener_report.md"

    lines = [f"# Screener Report — {date.today().strftime('%d/%m/%Y')}\n"]
    lines.append(f"Candidatas encontradas: {screener_result.get('total_candidates_found', 0)}\n")

    lines.append("## Top 5 candidatas\n")
    for item in screener_result.get("top_5", []):
        lines.append(f"### #{item['rank']} — {item['ticker']}: {item.get('name', '')}")
        lines.append(f"{item.get('reason', '')}\n")

    if screener_result.get("discarded"):
        lines.append("## Descartadas\n")
        for d in screener_result["discarded"]:
            lines.append(f"- {d}")

    path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  [doc] Screener report guardado: {path}")
    return path


# ─── Helpers DOCX ───────────────────────────────────────────────────────────────

def _create_styled_document() -> Document:
    """Crea un documento Word con estilos base."""
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Estilo de párrafo normal
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    return doc


def _add_cover(doc: Document, ticker: str, company_name: str, date_str: str):
    """Añade página de portada."""
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"TESIS DE INVERSIÓN")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(f"{ticker}" + (f" — {company_name}" if company_name else ""))
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(date.today().strftime("%d de %B de %Y")).font.size = Pt(12)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = disclaimer.add_run("Documento de análisis personal — No es consejo de inversión")
    run3.italic = True
    run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()


def _markdown_to_docx(doc: Document, text: str):
    """Convierte Markdown básico a párrafos DOCX con formato."""
    for line in text.split("\n"):
        stripped = line.strip()

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.", stripped):
            text_content = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(text_content, style="List Number")
        elif stripped == "" or stripped == "---":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            # Procesar bold e italic inline
            _add_formatted_run(p, stripped)


def _add_formatted_run(paragraph, text: str):
    """Aplica formato bold/italic dentro de un párrafo."""
    # Patrón: **bold**, *italic*, `code`
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)"
    for match in re.finditer(pattern, text):
        chunk = match.group()
        run = paragraph.add_run()
        if chunk.startswith("**") and chunk.endswith("**"):
            run.text = chunk[2:-2]
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run.text = chunk[1:-1]
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run.text = chunk[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run.text = chunk


def _add_footer(doc: Document):
    """Añade disclaimer al pie del último párrafo."""
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "Este documento es de uso personal y educativo. "
        "No constituye consejo de inversión. "
        "Invierte con criterio propio y conocimiento de los riesgos."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
